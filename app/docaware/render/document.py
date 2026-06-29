"""docaware/render/document.py — Write formatted Markdown to downloadable files.

Formats:
* ``md``   — always available, zero dependencies (formulas stay as LaTeX).
* ``docx`` — python-docx; block formulas ($$...$$) become embedded images.
* ``pdf``  — Pandoc + LaTeX if present (best fidelity), else an fpdf2 fallback.

The goal is a file the user can download/keep. Each backend degrades gracefully
and reports clearly if its optional dependency is missing.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from ..errors import BackendNotInstalledError

SUPPORTED_FORMATS = ("md", "docx", "pdf")

_BLOCK_FORMULA = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")


def render_document(markdown: str, out_path: str | Path, fmt: str = "md") -> Path:
    """Render ``markdown`` to ``out_path`` in the requested format.

    Args:
        markdown: GitHub-flavored Markdown (may contain LaTeX in $...$/$$...$$).
        out_path: Destination path (suffix is normalized to match ``fmt``).
        fmt: One of ``SUPPORTED_FORMATS``.

    Returns:
        The path actually written.

    Raises:
        ValueError: For an unsupported format.
    """
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format {fmt!r}; choose from {SUPPORTED_FORMATS}")
    out_path = Path(out_path).with_suffix(f".{fmt}")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt == "md":
        out_path.write_text(markdown, encoding="utf-8")
    elif fmt == "docx":
        _render_docx(markdown, out_path)
    elif fmt == "pdf":
        _render_pdf(markdown, out_path)
    return out_path


def _render_docx(markdown: str, out_path: Path) -> None:
    """Minimal Markdown → DOCX with embedded images for block formulas."""
    try:
        import docx  # type: ignore
        from docx.shared import Inches
    except ImportError as exc:
        raise BackendNotInstalledError(
            "python-docx not installed: pip install python-docx"
        ) from exc

    document = docx.Document()
    img_dir = Path(tempfile.mkdtemp(prefix="adtc_formula_"))
    counter = 0

    for raw_line in markdown.split("\n"):
        line = raw_line.rstrip()
        if not line:
            continue

        block = _BLOCK_FORMULA.search(line)
        if block:
            # Render the equation to an image and insert it centered.
            from .formula_img import render_latex_png

            counter += 1
            png = render_latex_png(block.group(1).strip(), img_dir / f"f{counter}.png")
            document.add_picture(str(png), width=Inches(3.0))
            continue

        heading = _HEADING.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2), level=level)
            continue

        bullet = _BULLET.match(line)
        if bullet:
            document.add_paragraph(bullet.group(1), style="List Bullet")
            continue

        document.add_paragraph(line)

    document.save(str(out_path))


def _render_pdf(markdown: str, out_path: Path) -> None:
    """Render to PDF via Pandoc+LaTeX if available, else an fpdf2 fallback."""
    if shutil.which("pandoc"):
        # Best fidelity: Pandoc typesets LaTeX math properly.
        with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
            tmp.write(markdown)
            md_path = tmp.name
        try:
            subprocess.run(
                ["pandoc", md_path, "-o", str(out_path)],
                check=True,
                capture_output=True,
            )
            return
        except subprocess.CalledProcessError:
            pass  # fall through to the lightweight fallback
        finally:
            Path(md_path).unlink(missing_ok=True)

    _render_pdf_fallback(markdown, out_path)


_FONT_DIR = Path(__file__).parent / "fonts"


def _render_pdf_fallback(markdown: str, out_path: Path) -> None:
    """Dependency-light PDF via fpdf2 with a bundled Unicode font (DejaVuSans).

    The Unicode font renders any character (em-dash, ω, π, √, …) so the PDF never
    crashes on the model's output; ``wrapmode="CHAR"`` guards against unbreakable
    long tokens. Headings/bullets are styled; math stays as readable LaTeX text.
    """
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError as exc:
        raise BackendNotInstalledError(
            "No PDF backend available. Install Pandoc (recommended) or fpdf2:\n"
            "  pip install fpdf2"
        ) from exc

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    regular, bold = _FONT_DIR / "DejaVuSans.ttf", _FONT_DIR / "DejaVuSans-Bold.ttf"
    if regular.exists():  # Unicode-capable (preferred)
        pdf.add_font("dejavu", "", str(regular))
        if bold.exists():
            pdf.add_font("dejavu", "B", str(bold))
        family, uni = "dejavu", True
    else:  # last-resort core font (latin-1 only)
        family, uni = "Helvetica", False

    def write(text: str, *, size: int, style: str = "") -> None:
        pdf.set_font(family, style=style if (uni or style == "") else style, size=size)
        if not uni:
            text = text.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, size * 0.55 + 2, text, wrapmode="CHAR")

    for raw in markdown.split("\n"):
        line = raw.rstrip()
        if not line:
            pdf.ln(3)
            continue
        heading = _HEADING.match(line)
        bullet = _BULLET.match(line)
        if heading:
            write(heading.group(2), size=15, style="B")
        elif bullet:
            write("•  " + bullet.group(1), size=11)
        else:
            write(line, size=11)
    pdf.output(str(out_path))
